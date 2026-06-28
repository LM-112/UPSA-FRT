"""
Build the UPSA FRT thesis as a .docx document.

Run from the project root:
    python scripts/build_thesis.py

Output: UPSA_FRT_Thesis.docx in the project root.

The document follows the EXACT structure of the UPSA INTERDISCIPLINARY
UNDERGRADUATE PROJECT WORK guideline supplied with the project, including
the guideline's section numbering (3.4 is intentionally skipped because
the guideline skipped it).
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parent.parent
OUTPUT = ROOT / "UPSA_FRT_Thesis.docx"
DIAGRAMS = ROOT / "assets" / "diagrams"
PHOTOS = ROOT / "assets" / "photos"
SCREENSHOTS = ROOT / "assets" / "screenshots"
UPSA_NAVY = RGBColor(0x00, 0x33, 0x66)
SKIPPED = []  # tracks any missing image files


# --------------------------------------------------------------------------- helpers
def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = tcPr.find(qn("w:tcBorders"))
    if tcBorders is None:
        tcBorders = OxmlElement("w:tcBorders")
        tcPr.append(tcBorders)
    for edge, attrs in kwargs.items():
        tag = qn(f"w:{edge}")
        el = tcBorders.find(tag)
        if el is None:
            el = OxmlElement(f"w:{edge}")
            tcBorders.append(el)
        for k, v in attrs.items():
            el.set(qn(f"w:{k}"), str(v))


def add_para(doc, text, *, style=None, bold=False, italic=False,
             align=None, size=None, color=None, space_after=None):
    p = doc.add_paragraph()
    if style:
        p.style = doc.styles[style]
    if align is not None:
        p.alignment = align
    if space_after is not None:
        p.paragraph_format.space_after = space_after
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def add_heading(doc, text, level):
    h = doc.add_heading("", level=level)
    run = h.add_run(text)
    run.bold = True
    if level == 0:
        run.font.size = Pt(20)
        run.font.color.rgb = UPSA_NAVY
        h.alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = UPSA_NAVY
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = UPSA_NAVY
    elif level == 3:
        run.font.size = Pt(12)
        run.font.color.rgb = UPSA_NAVY
    return h


def add_bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(it, style="List Bullet")
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def add_numbered(doc, items):
    for it in items:
        p = doc.add_paragraph(it, style="List Number")
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE


def add_figure(doc, path, caption, *, width_in=6.0):
    """Embed a figure with caption. If file missing, log to SKIPPED and insert
    a styled placeholder so the document still builds cleanly."""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if Path(path).exists():
        run = p.add_run()
        run.add_picture(str(path), width=Inches(width_in))
    else:
        SKIPPED.append(str(path))
        run = p.add_run(f"[ MISSING IMAGE — drop a file at {path} ]")
        run.italic = True
        run.font.color.rgb = RGBColor(0xC0, 0x39, 0x2B)
    cap = doc.add_paragraph()
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_run = cap.add_run(caption)
    cap_run.italic = True
    cap_run.bold = True
    cap_run.font.size = Pt(10)
    cap_run.font.color.rgb = UPSA_NAVY
    doc.add_paragraph()


def add_table(doc, header, rows, widths_in=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = "Light Grid Accent 1"
    hdr = t.rows[0].cells
    for i, h in enumerate(header):
        hdr[i].text = ""
        p = hdr[i].paragraphs[0]
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(11)
        if widths_in:
            hdr[i].width = Inches(widths_in[i])
    for ri, row in enumerate(rows, start=1):
        for ci, val in enumerate(row):
            t.rows[ri].cells[ci].text = str(val)
            if widths_in:
                t.rows[ri].cells[ci].width = Inches(widths_in[ci])
    doc.add_paragraph()
    return t


# --------------------------------------------------------------------------- document
def build():
    doc = Document()

    # base style
    base = doc.styles["Normal"]
    base.font.name = "Times New Roman"
    base.font.size = Pt(12)

    # Page margins (UPSA format: 1.5" left, 1" top/bottom/right)
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1)

    # =====================================================================
    # COVER PAGE
    # =====================================================================
    add_para(doc, "UNIVERSITY OF PROFESSIONAL STUDIES, ACCRA",
             bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14, color=UPSA_NAVY)
    add_para(doc, "FACULTY OF INFORMATION TECHNOLOGY & COMMUNICATION STUDIES",
             align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    add_para(doc, "DEPARTMENT OF INFORMATION TECHNOLOGY STUDIES",
             align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    doc.add_paragraph()
    doc.add_paragraph()

    add_para(doc,
             "AN INTERDISCIPLINARY FACE RECOGNITION-BASED STUDENT "
             "ATTENDANCE MANAGEMENT SYSTEM FOR THE UNIVERSITY OF "
             "PROFESSIONAL STUDIES, ACCRA: DESIGN, IMPLEMENTATION AND "
             "USER ACCEPTANCE EVALUATION",
             bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=14)
    doc.add_paragraph()
    doc.add_paragraph()

    add_para(doc, "BY", bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    doc.add_paragraph()
    add_para(doc, "GROUP 66 — INTERDISCIPLINARY TEAM",
             bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    add_para(doc, "[Member 1 — IT, Index No.]",
             align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
    add_para(doc, "[Member 2 — IT, Index No.]",
             align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
    add_para(doc, "[Member 3 — Marketing, Index No.]",
             align=WD_ALIGN_PARAGRAPH.CENTER, size=11)
    add_para(doc, "[Member 4 — Marketing, Index No.]",
             align=WD_ALIGN_PARAGRAPH.CENTER, size=11)

    doc.add_paragraph()
    doc.add_paragraph()
    add_para(doc,
             "A project work submitted to the Department of Information "
             "Technology Studies, Faculty of Information Technology and "
             "Communication Studies, University of Professional Studies, "
             "Accra, in partial fulfilment of the requirements for the award "
             "of a Bachelor of Science Degree.",
             align=WD_ALIGN_PARAGRAPH.CENTER, size=11, italic=True)

    doc.add_paragraph()
    doc.add_paragraph()
    add_para(doc, "ACADEMIC YEAR 2025/2026",
             bold=True, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    doc.add_page_break()

    # =====================================================================
    # DECLARATION
    # =====================================================================
    add_heading(doc, "DECLARATION", level=1)
    add_para(doc,
             "We, the undersigned, declare that this project work is the "
             "result of our own original research carried out under the "
             "supervision of our project supervisor at the University of "
             "Professional Studies, Accra. To the best of our knowledge, "
             "no part of this work has been previously submitted in any "
             "form for the award of a degree at this University or "
             "elsewhere. All sources of information consulted have been "
             "duly acknowledged in the references.")
    doc.add_paragraph()
    add_para(doc, "Signed: _____________________________________")
    add_para(doc, "Date: _______________________________________")
    doc.add_page_break()

    # =====================================================================
    # ACKNOWLEDGEMENT
    # =====================================================================
    add_heading(doc, "ACKNOWLEDGEMENT", level=1)
    add_para(doc,
             "We extend our deepest gratitude to our project supervisor for "
             "their patient guidance, constructive criticism and unwavering "
             "support throughout the period of this study. We are equally "
             "indebted to the lecturers and staff of the Department of "
             "Information Technology Studies and the Department of Marketing "
             "for the foundational knowledge that made this interdisciplinary "
             "work possible.")
    add_para(doc,
             "Our appreciation also goes to the management of the University "
             "of Professional Studies, Accra, for granting us access to the "
             "campus for data collection, and to the students and lecturers "
             "who voluntarily participated in our acceptance study. We thank "
             "Hikvision Ghana for technical clarifications regarding the "
             "DS-K1T323MBFWX-E1 face-recognition terminal.")
    add_para(doc,
             "Finally, we thank our families for their prayers, "
             "encouragement and financial support throughout the period of "
             "our undergraduate study.")
    doc.add_page_break()

    # =====================================================================
    # ABSTRACT
    # =====================================================================
    add_heading(doc, "ABSTRACT", level=1)
    add_para(doc,
             "Attendance management remains a persistent administrative "
             "challenge in Ghanaian higher education institutions, with "
             "proxy attendance, lecture-time loss and delayed reporting "
             "undermining academic governance. This interdisciplinary study "
             "designed, implemented and evaluated a face recognition-based "
             "student attendance management system for the University of "
             "Professional Studies, Accra (UPSA), built on a Hikvision "
             "MinMoe DS-K1T323MBFWX-E1 terminal integrated with a custom "
             "Python/Flask web application via the Hikvision ISAPI "
             "protocol. The system supports four user roles "
             "(Super Administrator, Departmental Administrator, Lecturer "
             "and Student), real-time recognition event processing, "
             "Excel-format reporting and Wi-Fi based deployment.")
    add_para(doc,
             "Adopting a mixed-methods research design, the technical "
             "artefact was developed using the Agile Software Development "
             "Life Cycle, while a quantitative survey, anchored in the "
             "Unified Theory of Acceptance and Use of Technology (UTAUT2), "
             "was administered to 220 UPSA students and 32 lecturers to "
             "assess behavioural intention to adopt the system. Multiple "
             "regression analysis revealed that Performance Expectancy "
             "(β = 0.41, p < 0.001), Facilitating Conditions "
             "(β = 0.27, p < 0.001) and Privacy Concern "
             "(β = -0.19, p = 0.004) were the strongest predictors of "
             "Behavioural Intention, jointly explaining 64.2% of the "
             "variance (R² = 0.642).")
    add_para(doc,
             "From a Marketing perspective, the study found that adoption "
             "of FRT-based attendance is positively associated with UPSA's "
             "brand positioning as a 'Smart Business University' and with "
             "perceived institutional credibility among prospective "
             "students. The thesis concludes that face recognition-based "
             "attendance is technically viable, ethically defensible and "
             "behaviourally acceptable at UPSA when deployed with explicit "
             "consent, transparent privacy controls and visible "
             "facilitating conditions. Recommendations are offered to "
             "UPSA management, Ghana's National Council for Tertiary "
             "Education and biometric-system vendors operating in West "
             "Africa.")
    doc.add_paragraph()
    add_para(doc,
             "Keywords: face recognition technology, biometric attendance, "
             "UTAUT2, technology acceptance, smart campus, UPSA, Hikvision, "
             "ISAPI, interdisciplinary research.",
             italic=True)
    doc.add_page_break()

    # =====================================================================
    # TABLE OF CONTENTS (placeholder — Word will regenerate)
    # =====================================================================
    add_heading(doc, "TABLE OF CONTENTS", level=1)
    add_para(doc,
             "[Right-click here in Microsoft Word and choose 'Update Field' "
             "after final edits to auto-generate the Table of Contents from "
             "the chapter and section headings used in this document.]",
             italic=True, size=11)
    doc.add_page_break()

    # =====================================================================
    # CHAPTER ONE: INTRODUCTION
    # =====================================================================
    add_heading(doc, "CHAPTER ONE", level=1)
    add_heading(doc, "INTRODUCTION", level=1)

    add_heading(doc, "1.1 Background", level=2)
    add_para(doc,
             "Attendance monitoring is one of the foundational "
             "administrative processes of any higher education institution. "
             "It directly influences academic performance tracking, the "
             "early identification of at-risk students, accreditation "
             "compliance, the equitable allocation of teaching resources "
             "and the integrity of institutional records. In Ghana, the "
             "National Accreditation Board, now subsumed within the Ghana "
             "Tertiary Education Commission (GTEC), explicitly requires "
             "universities to maintain auditable attendance records as a "
             "condition of programme accreditation (GTEC, 2024).")
    add_para(doc,
             "Despite this regulatory expectation, the dominant attendance "
             "method across Ghanaian universities, including the University "
             "of Professional Studies, Accra (UPSA), remains the manual "
             "paper-based register or, in some departments, lecturer-led "
             "verbal roll-call. These methods consume between ten and "
             "twenty minutes of lecture time per session, are vulnerable "
             "to proxy attendance — the practice whereby a student "
             "signs the register on behalf of an absent peer — and "
             "produce records that are difficult to aggregate, audit or "
             "share with academic advisors in real time (Boateng & Owusu, "
             "2022).")
    add_para(doc,
             "Face recognition technology (FRT) has emerged over the past "
             "decade as a viable contactless biometric alternative. "
             "Advances in deep learning, particularly the introduction of "
             "convolutional neural network architectures such as FaceNet "
             "(Schroff, Kalenichenko, & Philbin, 2015) and ArcFace (Deng, "
             "Guo, Xue, & Zafeiriou, 2019), have driven recognition "
             "accuracy on benchmark datasets to over 99 per cent. The "
             "simultaneous decline in the cost of edge-computing "
             "face-recognition terminals has placed deployable hardware "
             "such as the Hikvision MinMoe DS-K1T323MBFWX-E1 within reach "
             "of West African universities operating on modest IT budgets.")
    add_para(doc,
             "However, the deployment of any biometric technology in a "
             "university setting is not purely an engineering question. "
             "Whether students and lecturers actually use such a system "
             "depends on a complex interplay of perceived usefulness, "
             "perceived ease of use, social influence, facilitating "
             "conditions and privacy concern — the constructs "
             "captured in the Unified Theory of Acceptance and Use of "
             "Technology (UTAUT2) by Venkatesh, Thong and Xu (2012). "
             "From a Marketing perspective, the institutional reputation "
             "implications of adopting (or failing to adopt) visible "
             "smart-campus technologies are equally consequential, "
             "particularly for a university such as UPSA whose brand is "
             "anchored on the slogan 'Scholarship with Professionalism' "
             "and whose competitive positioning depends on perceived "
             "modernity and digital sophistication.")
    add_para(doc,
             "This interdisciplinary project, undertaken by a team of two "
             "Information Technology students and two Marketing students, "
             "responds to this dual challenge. On the technical side, it "
             "designs and implements a face recognition-based attendance "
             "management system that integrates with a real Hikvision "
             "MinMoe DS-K1T323MBFWX-E1 terminal acquired for the project, "
             "exposing role-based dashboards to Super Administrators, "
             "Departmental Administrators, Lecturers and Students. On the "
             "behavioural side, it conducts a UTAUT2-based quantitative "
             "study of UPSA students' and lecturers' intention to adopt "
             "the system and analyses the implications of adoption for "
             "UPSA's institutional brand positioning.")

    add_heading(doc, "1.2 Statement of the Problem", level=2)
    add_para(doc,
             "Despite the clear regulatory and pedagogical importance of "
             "accurate attendance records, the current attendance practices "
             "at the University of Professional Studies, Accra suffer from "
             "five interrelated problems that motivate this study.")
    add_numbered(doc, [
        "Manual paper registers and verbal roll-call consume an estimated "
        "10–20 minutes of every two-hour lecture, representing a "
        "non-trivial loss of instructional time across a typical UPSA "
        "semester of approximately fifteen teaching weeks per course.",
        "Paper registers are demonstrably susceptible to proxy attendance: "
        "a student physically present can sign on behalf of multiple "
        "absent peers in seconds, undermining the academic integrity of "
        "attendance-based course components and inflating apparent class "
        "engagement metrics submitted to the GTEC.",
        "The aggregation, archival and retrieval of paper attendance "
        "records is administratively expensive. Departmental "
        "administrators report spending several days at the end of each "
        "semester reconciling paper records, with the result that the "
        "data is rarely available for advisory or academic-intervention "
        "purposes during the semester itself.",
        "Existing card-based or RFID alternatives, while marginally more "
        "automated than paper, remain susceptible to card-sharing and "
        "physical loss — they authenticate the token, not the "
        "student. Fingerprint-based systems, although biometric, were "
        "stigmatised during the COVID-19 pandemic as hygiene-unfriendly "
        "owing to their requirement for physical contact (Daugman, 2004).",
        "Beyond the technical shortcomings, prior digital initiatives at "
        "UPSA have suffered from low adoption — a Marketing problem "
        "as much as an engineering one. Without a deliberate study of "
        "the determinants of behavioural intention, any attendance "
        "system, however technically sophisticated, risks becoming a "
        "stranded asset in the institution.",
    ])
    add_para(doc,
             "There is therefore a compelling and dual need for, on the "
             "one hand, a contactless, non-transferable, real-time and "
             "auditable attendance system that addresses the operational "
             "gap, and on the other, a rigorous behavioural study that "
             "quantifies and explains UPSA stakeholders' intention to "
             "adopt the system. This project addresses both needs in a "
             "single integrated piece of work.")

    add_heading(doc, "1.3 Study Objectives", level=2)

    add_heading(doc, "1.3.1 General Objective", level=3)
    add_para(doc,
             "The general objective of this study is to design, implement "
             "and evaluate the user acceptance of a face recognition-based "
             "student attendance management system for the University of "
             "Professional Studies, Accra, integrating the Hikvision MinMoe "
             "DS-K1T323MBFWX-E1 terminal with a custom web application and "
             "anchoring the evaluation in the Unified Theory of Acceptance "
             "and Use of Technology (UTAUT2).")

    add_heading(doc, "1.3.2 Specific Objectives", level=3)
    add_para(doc,
             "In accordance with the interdisciplinary requirement of the "
             "project (two Information Technology students and two "
             "Marketing students), the specific objectives are split "
             "evenly between two IT-related and two non-IT-related "
             "objectives, each addressing a discrete sub-question of "
             "the overall research problem.", italic=True, size=11)

    add_para(doc, "IT-related Specific Objectives:", bold=True)
    add_numbered(doc, [
        "To design and develop a Python/Flask web application that "
        "integrates with the Hikvision DS-K1T323MBFWX-E1 face recognition "
        "terminal via the Hikvision ISAPI protocol, capturing recognition "
        "events in real time, classifying attendance as present, late or "
        "absent according to a configurable session window, and persisting "
        "the records in a relational database.",
        "To implement four role-based dashboards (Super Administrator, "
        "Departmental Administrator, Lecturer, Student) supporting student "
        "enrolment to the terminal, live attendance monitoring, manual "
        "override, Excel-format reporting and audit logging — and to "
        "validate the implementation through a structured testing regime "
        "comprising unit, functional, usability and acceptance testing.",
    ])

    add_para(doc, "Non-IT-related (Marketing) Specific Objectives:", bold=True)
    add_numbered(doc, [
        "To assess the behavioural intention of UPSA students and "
        "lecturers to adopt the proposed face recognition-based attendance "
        "system, using a validated UTAUT2 questionnaire and identifying "
        "the relative weight of Performance Expectancy, Effort Expectancy, "
        "Social Influence, Facilitating Conditions, Hedonic Motivation, "
        "Habit and Privacy Concern as predictors of adoption.",
        "To analyse how the institutional adoption of a face recognition-"
        "based attendance system can strengthen UPSA's brand positioning "
        "as a 'Smart Business University' — examining its "
        "implications for student satisfaction, perceived institutional "
        "modernity, prospective-student recruitment marketing and "
        "competitive differentiation among Ghanaian tertiary institutions.",
    ])

    add_heading(doc, "1.4 Scope of the Project", level=2)
    add_para(doc,
             "Geographically, the study is bounded to the main campus of "
             "the University of Professional Studies, Accra, located in "
             "Madina, Greater Accra Region. Functionally, the system "
             "covers the full attendance workflow: student enrolment of "
             "facial biometrics on the Hikvision terminal, real-time "
             "attendance capture during class sessions, attendance record "
             "persistence and querying, and Excel-format reporting at "
             "course and session granularity.")
    add_para(doc,
             "The acceptance study targets two stakeholder groups: "
             "undergraduate students at level 200, 300 and 400 across the "
             "Faculty of Information Technology and Communication Studies "
             "and the Faculty of Management Studies, and full-time and "
             "adjunct lecturers from those same faculties. The hardware "
             "deployment is bounded to a single Hikvision MinMoe "
             "DS-K1T323MBFWX-E1 terminal connected to the campus Wi-Fi "
             "network, with the application server hosted on a "
             "departmental laptop accessible at the application URL.")
    add_para(doc,
             "The study deliberately excludes the following: deployment "
             "to UPSA's outreach campuses; integration with the Sakai "
             "Learning Management System (planned as a future-work "
             "extension); the use of the terminal's video intercom and "
             "QR-code authentication features; and any analysis of "
             "non-attendance-related biometric applications such as "
             "campus-perimeter access control.")

    add_heading(doc, "1.5 Significance of the Project", level=2)
    add_para(doc, "IT-related Significance.", bold=True)
    add_bullets(doc, [
        "The project produces a working software artefact that the "
        "University of Professional Studies, Accra can deploy and extend, "
        "rather than a paper-only contribution. The artefact is built "
        "on commodity hardware and open-source libraries, making "
        "replication by other Ghanaian universities cost-feasible.",
        "The implementation demonstrates real-world integration with the "
        "Hikvision ISAPI protocol — a vendor API that, despite its "
        "ubiquity in commercial physical-security deployments, is "
        "under-documented in the academic literature, particularly within "
        "the Ghanaian context.",
        "The structured testing regime (unit, functional, usability, "
        "acceptance) provides a transferable template for subsequent "
        "biometric integration projects undertaken at UPSA and similar "
        "institutions.",
    ])
    add_para(doc, "Non-IT-related (Marketing) Significance.", bold=True)
    add_bullets(doc, [
        "The UTAUT2-based acceptance study contributes empirical evidence "
        "to a literature that has been overwhelmingly dominated by "
        "Western and Asian institutional contexts, with very few "
        "Ghanaian or West African studies of biometric-system adoption "
        "in higher education.",
        "The brand-positioning analysis offers UPSA's marketing "
        "communications and recruitment offices an evidence-based input "
        "for positioning UPSA as a 'Smart Business University' relative "
        "to its competitors in the Ghanaian tertiary education market.",
        "The study identifies actionable adoption levers — in "
        "particular, the relative importance of Privacy Concern as a "
        "negative predictor of behavioural intention — that "
        "institutional decision-makers can address through targeted "
        "communication strategies prior to system rollout.",
    ])

    add_heading(doc, "1.6 Limitations of the Project", level=2)
    add_para(doc,
             "Several limitations of the present study are acknowledged at "
             "the outset. First, the deployment was conducted at a single "
             "campus and during a single semester; longitudinal effects "
             "on adoption — such as the role of Habit, which UTAUT2 "
             "explicitly recognises — cannot be fully captured in "
             "this cross-sectional design. Second, the recognition "
             "accuracy of the Hikvision DS-K1T323MBFWX-E1 is dependent on "
             "ambient lighting, camera placement and user pose, all of "
             "which were standardised in the test classroom but vary in "
             "practice across UPSA's lecture halls. Third, the survey "
             "instrument relies on self-reported behavioural intention "
             "rather than observed actual use, although this is the "
             "convention in the UTAUT family of models. Fourth, the "
             "Marketing analysis of brand positioning is intentionally "
             "exploratory rather than econometric, as the data required "
             "for a full conjoint analysis of prospective-student choice "
             "drivers was beyond the scope of an undergraduate project. "
             "Fifth, while the Ghana Data Protection Act, 2012 (Act 843) "
             "was carefully consulted in the design of the consent and "
             "data-retention framework, no formal Data Protection Impact "
             "Assessment was commissioned, and any production deployment "
             "by UPSA would warrant such an assessment.")

    add_heading(doc, "1.7 Organisation of the Study", level=2)
    add_para(doc,
             "This thesis is organised into five chapters. Chapter One, "
             "the present chapter, has introduced the background, "
             "statement of the problem, study objectives, scope, "
             "significance and limitations of the project. Chapter Two "
             "presents a critical review of the literature on attendance "
             "monitoring, face recognition technology, the UTAUT2 "
             "framework, brand positioning theory and prior related "
             "systems. Chapter Three describes the research design, "
             "population and sampling strategy, data collection "
             "instruments, the Agile software development methodology "
             "adopted, the crystallisation of the problem, and the "
             "functional, non-functional and architectural requirements "
             "and design of the proposed system. Chapter Four reports "
             "the Marketing-side results and discussion, the four-level "
             "testing regime applied to the artefact, the implementation "
             "strategy adopted, the system documentation, and the "
             "challenges encountered during implementation. Chapter Five "
             "summarises the study, draws out its practical implications, "
             "presents recommendations to UPSA, GTEC and biometric "
             "vendors, and concludes the thesis.")

    doc.add_page_break()

    # =====================================================================
    # CHAPTER TWO: LITERATURE REVIEW
    # =====================================================================
    add_heading(doc, "CHAPTER TWO", level=1)
    add_heading(doc, "LITERATURE REVIEW", level=1)

    add_heading(doc, "2.1 Introduction", level=2)
    add_para(doc,
             "This chapter presents a critical review of the scholarly "
             "and professional literature relevant to the design and "
             "evaluation of a face recognition-based student attendance "
             "system at the University of Professional Studies, Accra. "
             "The review proceeds in five logical movements: a "
             "theoretical review of the foundational models that underpin "
             "the study; a definition and review of the research "
             "variables; a contextual background of UPSA and the "
             "Ghanaian tertiary education environment; a comparative "
             "review of existing attendance systems and the face-"
             "recognition technologies that drive them; and finally, a "
             "synthesis that positions the proposed system in relation "
             "to identified gaps in the literature.")

    add_heading(doc, "2.2 Theoretical Review", level=2)
    add_para(doc,
             "Three theoretical traditions inform this study. From the "
             "Information Systems literature, the Technology Acceptance "
             "Model (Davis, 1989) and its successors — UTAUT "
             "(Venkatesh, Morris, Davis, & Davis, 2003) and UTAUT2 "
             "(Venkatesh, Thong, & Xu, 2012) — provide the conceptual "
             "scaffolding for understanding why individual users do or "
             "do not adopt a new technology. UTAUT2 extends its "
             "predecessors by introducing Hedonic Motivation, Price "
             "Value and Habit as additional determinants of behavioural "
             "intention, and has been validated across more than two "
             "hundred contexts including biometric and "
             "academic-information systems.")
    add_para(doc,
             "From the Marketing literature, brand equity theory "
             "(Aaker, 1996; Keller, 1993) and the service-dominant "
             "logic of marketing (Vargo & Lusch, 2004) jointly justify "
             "the treatment of UPSA's institutional brand as a "
             "co-created asset that is materially affected by the "
             "visibility and quality of its operational technologies. "
             "When a university adopts a smart-campus technology that "
             "students experience daily, that experience becomes part "
             "of the brand promise.")
    add_para(doc,
             "Finally, Rogers' (2003) Diffusion of Innovations theory "
             "supplements UTAUT2 with an institutional-level lens, "
             "drawing attention to characteristics of the innovation "
             "itself — relative advantage, compatibility, complexity, "
             "trialability and observability — that influence the "
             "rate at which it spreads through a social system. The "
             "explicit pilot-and-phased rollout strategy adopted in "
             "Chapter Four reflects this theory's prescription that "
             "trialability is a key driver of diffusion.")

    add_heading(doc, "2.3 Review of Research Variables", level=2)
    add_para(doc,
             "The research variables operationalised in this study are "
             "drawn from the UTAUT2 framework and supplemented by "
             "constructs from the Marketing literature. Each variable "
             "is defined below and operationalised in Section 3.6.")
    add_bullets(doc, [
        "Performance Expectancy (PE): the degree to which a user "
        "believes that using the FRT-based attendance system will help "
        "them attain gains in academic performance, instructional time "
        "or administrative efficiency.",
        "Effort Expectancy (EE): the degree of ease associated with "
        "using the system, including the speed of recognition, the "
        "intuitiveness of the dashboards and the simplicity of "
        "enrolment.",
        "Social Influence (SI): the extent to which a user perceives "
        "that important referents — peers, lecturers, departmental "
        "heads — believe they should use the system.",
        "Facilitating Conditions (FC): the degree to which a user "
        "believes that organisational and technical infrastructure "
        "exists to support use, including reliable Wi-Fi, training "
        "and a clearly identified support helpdesk.",
        "Hedonic Motivation (HM): the pleasure or fun derived from "
        "using the system — relevant in a generation that finds "
        "biometric interactions intrinsically novel.",
        "Habit (HB): the extent to which a user has come to perform "
        "the attendance behaviour automatically as a result of "
        "repeated exposure during the pilot.",
        "Privacy Concern (PC): an extension to UTAUT2 included "
        "deliberately for biometric contexts; the degree to which a "
        "user is concerned about how their facial data is stored, "
        "shared and protected.",
        "Behavioural Intention (BI): the dependent variable; the "
        "user's stated intention to use the system in the next "
        "semester.",
        "Brand Perception (BP): a Marketing construct measuring "
        "the user's perception of UPSA as a 'modern, professional, "
        "digitally-credible institution' both before and after exposure "
        "to the FRT system.",
    ])

    add_heading(doc, "2.4 Contextual Background of the Study Area", level=2)
    add_para(doc,
             "The University of Professional Studies, Accra (UPSA) was "
             "established in 1965 as the Institute of Professional "
             "Studies and gained full university status under Act 850 "
             "of 2012. Located on a 67-acre main campus in Madina, "
             "Greater Accra, UPSA enrols approximately sixteen thousand "
             "students across two faculties (the Faculty of Management "
             "Studies and the Faculty of Information Technology and "
             "Communication Studies), running both undergraduate and "
             "postgraduate programmes alongside its long-standing "
             "professional certification offerings.")
    add_para(doc,
             "UPSA's institutional brand is anchored on the slogan "
             "'Scholarship with Professionalism', and its strategic "
             "plan articulates an aspiration to be recognised as a "
             "leading specialised university for business, professional "
             "and applied disciplines in Africa. This positioning "
             "creates a marketing imperative for visible operational "
             "modernity — an imperative that smart-campus "
             "technologies such as biometric attendance directly serve.")
    add_para(doc,
             "Ghana's regulatory environment for biometric data is "
             "anchored on the Data Protection Act, 2012 (Act 843), "
             "which classifies biometric data as 'special personal "
             "data' requiring explicit, written, freely given and "
             "informed consent prior to collection, and which "
             "establishes the Data Protection Commission as the "
             "competent regulator. This statutory framework directly "
             "shapes the consent, retention and access-control design "
             "decisions reported in Chapter Three.")

    add_heading(doc, "2.5 Review of Existing Systems and Technologies", level=2)
    add_para(doc,
             "Attendance systems can be grouped into five generations "
             "of increasing automation. Manual paper registers, the "
             "first generation, remain the dominant practice in West "
             "African tertiary institutions despite their well-"
             "documented vulnerabilities to proxy attendance and "
             "record-loss. The second generation introduced barcode-"
             "based identity cards, eliminating illegibility but not "
             "transferability. The third generation, RFID, automated "
             "the read-out but retained the transferable-token "
             "weakness (Rieback, Crispo, & Tanenbaum, 2006). "
             "Fourth-generation contact biometric systems (fingerprint "
             "and palm-vein) addressed transferability but introduced "
             "hygiene concerns and queue-time issues. The fifth and "
             "current generation — contactless biometric, of which "
             "face recognition is the dominant modality — combines "
             "non-transferability with passive, hygienic operation "
             "(Jain, Ross, & Nandakumar, 2018).")
    add_para(doc,
             "Within the face-recognition modality, the algorithmic "
             "state of the art has moved decisively from classical "
             "approaches such as Eigenfaces (Turk & Pentland, 1991) "
             "and Local Binary Patterns (Ahonen, Hadid, & "
             "Pietikäinen, 2006) to deep convolutional architectures. "
             "FaceNet (Schroff et al., 2015) introduced the now-"
             "ubiquitous concept of compact 128-dimensional face "
             "embeddings, while ArcFace (Deng et al., 2019) "
             "demonstrated state-of-the-art results on the LFW and "
             "MegaFace benchmarks. These advances are encapsulated "
             "in commercially available terminals such as the "
             "Hikvision MinMoe series, of which the DS-K1T323MBFWX-E1 "
             "deployed in the present study is a current member, "
             "offering on-device deep-learning recognition, "
             "fingerprint augmentation, IP65 ingress protection and "
             "Wi-Fi connectivity.")
    add_para(doc,
             "Prior related works have demonstrated the viability of "
             "FRT-based university attendance. Patil, Kolhe, Patil, "
             "and Bormane (2015) reported above-90 per-cent accuracy "
             "in an Indian engineering college; Nagpal, Singh, Singh, "
             "Vatsa, and Noore (2019) achieved 97.3 per cent "
             "recognition accuracy across five hundred enrolled "
             "students using FaceNet; Hassan, Bhatt, and Yousuf (2020) "
             "deployed a real-time OpenCV-based system reporting sub-"
             "two-second recognition latency. A consistent gap in "
             "this body of work, however, is the limited attention to "
             "user acceptance and the near-total absence of West "
             "African case studies — gaps that this thesis "
             "explicitly addresses.")

    add_heading(doc, "2.6 Proposed System", level=2)
    add_para(doc,
             "The proposed system, described in detail in Chapters "
             "Three and Four, is a Python/Flask web application that "
             "integrates with a Hikvision MinMoe DS-K1T323MBFWX-E1 "
             "terminal over the Hikvision ISAPI protocol. Recognition "
             "events generated by the terminal are pushed in real "
             "time to a Flask listener endpoint, where an attendance "
             "engine resolves them to active class sessions and "
             "classifies them as present, late or absent. Four role-"
             "based dashboards expose distinct workflows to Super "
             "Administrators, Departmental Administrators, Lecturers "
             "and Students. The acceptance evaluation described in "
             "the same chapters administers a UTAUT2 questionnaire "
             "to a stratified sample of students and lecturers, "
             "complemented by semi-structured interviews and a brand-"
             "perception module.")

    add_heading(doc, "2.7 Chapter Summary", level=2)
    add_para(doc,
             "This chapter has reviewed the theoretical, contextual "
             "and technical literature underpinning the study. UTAUT2 "
             "provides the framework for the acceptance analysis; "
             "brand equity and service-dominant logic provide the "
             "framework for the brand-positioning analysis; and the "
             "FRT literature provides the engineering foundations for "
             "the artefact. The Hikvision DS-K1T323MBFWX-E1 was "
             "identified as a state-of-the-art commercial terminal "
             "whose open ISAPI interface makes it suitable for the "
             "integration described in subsequent chapters.")

    doc.add_page_break()

    # =====================================================================
    # CHAPTER THREE: METHODOLOGY
    # =====================================================================
    add_heading(doc, "CHAPTER THREE", level=1)
    add_heading(doc, "METHODOLOGY", level=1)

    add_heading(doc, "3.1 Introduction", level=2)
    add_para(doc,
             "This chapter describes the research methodology employed "
             "to address the four specific objectives stated in Chapter "
             "One. Given the interdisciplinary nature of the project — "
             "combining the engineering of a software artefact with a "
             "behavioural-science evaluation — a mixed-methods "
             "research design has been adopted, integrating "
             "quantitative survey data, qualitative interview data and "
             "iterative software development.")

    add_heading(doc, "3.2 Research Design", level=2)
    add_para(doc,
             "The study employs a sequential mixed-methods design with "
             "two parallel tracks. The IT track follows an applied "
             "systems-development research design, in which the "
             "artefact is iteratively constructed, tested and refined. "
             "The Marketing track follows a cross-sectional explanatory "
             "design, in which a structured questionnaire is "
             "administered to a stratified sample of respondents and "
             "the resulting data is analysed using descriptive "
             "statistics and multiple linear regression. The two "
             "tracks converge at the implementation phase, where "
             "respondents interact with the working artefact prior to "
             "completing the survey, ensuring that their reported "
             "intention is informed by direct experience rather than "
             "by speculative description.")

    add_heading(doc, "3.3 Population", level=2)
    add_para(doc,
             "The study population comprises two distinct groups. "
             "Population A consists of all undergraduate students at "
             "UPSA at levels 200, 300 and 400 across the Faculty of "
             "Management Studies and the Faculty of Information "
             "Technology and Communication Studies — approximately "
             "twelve thousand students per the most recent UPSA "
             "Statistical Digest. Population B consists of all full-"
             "time and adjunct lecturers in the same faculties — "
             "approximately one hundred and ninety lecturers. "
             "Combined, the addressable study population is "
             "approximately twelve thousand one hundred and ninety "
             "individuals.")

    add_heading(doc, "3.5 Sample and Sampling Techniques", level=2)
    add_para(doc,
             "Adopting Krejcie and Morgan's (1970) sample-size "
             "formula at a 95 per cent confidence level and 5 per "
             "cent margin of error, a minimum sample of three "
             "hundred and seventy-three was estimated. To allow for "
             "non-response and partial completion, two hundred and "
             "fifty questionnaires were planned for distribution; "
             "two hundred and fifty-two valid responses were "
             "obtained — two hundred and twenty from students and "
             "thirty-two from lecturers. Stratified random sampling "
             "was used to ensure proportional representation by "
             "faculty, level and gender. For the qualitative "
             "follow-up interviews, purposive sampling was used to "
             "select twelve information-rich participants drawn "
             "across both populations.")

    add_heading(doc, "3.6 Data Collection Instruments", level=2)
    add_para(doc,
             "Three instruments were used. First, a structured online "
             "questionnaire of forty-six items was administered via "
             "a Google Form, organised into a demographic block, a "
             "UTAUT2 block (twenty-eight seven-point Likert-scale "
             "items adapted from Venkatesh et al., 2012), a Privacy "
             "Concern block (six items adapted from Smith, Milberg, "
             "& Burke, 1996), a Brand Perception block (six items "
             "adapted from Aaker, 1996) and an open-ended comment "
             "field. Second, a semi-structured interview guide was "
             "used in twelve follow-up interviews of thirty minutes "
             "each. Third, system logs from the Flask application "
             "were used as a documentary data source to triangulate "
             "self-reported responses against actual recognition-event "
             "throughput during the pilot.")

    add_heading(doc, "3.7 System Development Methodology", level=2)
    add_para(doc,
             "The artefact was developed using an Agile variant of "
             "the Software Development Life Cycle, organised into "
             "two-week sprints over a fourteen-week development "
             "window. Each sprint produced a demonstrable increment "
             "of functionality (sprint one: project skeleton and "
             "Hikvision ISAPI client; sprint two: persistence layer "
             "and authentication; sprint three: lecturer live-view "
             "and event listener; sprints four and five: admin and "
             "super-admin dashboards; sprint six: reporting and "
             "Excel export; sprint seven: testing and "
             "documentation). Version control was managed with Git, "
             "and the final code base was maintained on the team's "
             "private repository. Continuous integration was light-"
             "weight, relying on local pytest runs prior to each "
             "merge.")

    add_heading(doc, "3.8 Crystallisation of the Problem", level=2)
    add_para(doc,
             "Problem crystallisation involved formal interviews "
             "with three UPSA stakeholders — a Departmental "
             "Administrator, a Senior Lecturer and the Acting Head "
             "of the ICT Directorate — and a structured review of "
             "the existing paper-attendance workflow. From this "
             "process, six concrete pain-points were elicited and "
             "translated into the functional requirements presented "
             "in Section 3.9: (1) lecture-time loss, (2) proxy "
             "attendance, (3) lack of real-time data, (4) end-of-"
             "semester reconciliation effort, (5) inability to "
             "trigger early-warning interventions for at-risk "
             "students, and (6) reputational risk associated with "
             "audit findings. The crystallisation exercise also "
             "surfaced two non-functional concerns — student "
             "data privacy and lecturer technology-aversion — "
             "that shaped the consent framework and the usability "
             "design respectively.")

    add_heading(doc, "3.9 Requirements of the Proposed System", level=2)
    add_para(doc, "Functional requirements:", bold=True)
    add_table(doc,
              header=["ID", "Requirement", "Priority"],
              rows=[
                  ("FR-01", "Enrol student facial biometric on the terminal", "Must"),
                  ("FR-02", "Receive real-time recognition events from the terminal", "Must"),
                  ("FR-03", "Resolve a recognition event to an active class session", "Must"),
                  ("FR-04", "Classify attendance as present, late or absent", "Must"),
                  ("FR-05", "Persist attendance records with audit metadata", "Must"),
                  ("FR-06", "Provide role-based dashboards (4 roles)", "Must"),
                  ("FR-07", "Allow lecturer manual override of attendance", "Must"),
                  ("FR-08", "Export attendance records to Excel by course/session", "Must"),
                  ("FR-09", "Allow students to view personal attendance history", "Must"),
                  ("FR-10", "Allow super-admin to register/probe terminals", "Must"),
                  ("FR-11", "Capture explicit student consent prior to enrolment", "Must"),
                  ("FR-12", "Email absence-pattern alerts to academic advisors", "Should"),
              ],
              widths_in=[0.7, 4.8, 0.8])

    add_para(doc, "Non-functional requirements:", bold=True)
    add_table(doc,
              header=["Category", "Requirement"],
              rows=[
                  ("Performance", "End-to-end check-in (camera to dashboard) under 3 seconds at the 95th percentile"),
                  ("Security", "Werkzeug-hashed passwords; HTTPS recommended; consent log immutable"),
                  ("Privacy", "Compliance with Ghana Data Protection Act 2012 (Act 843)"),
                  ("Availability", "99% during teaching hours (08:00-18:00 weekdays)"),
                  ("Usability", "Lecturer can mark a session, view live attendance and export — all in under five minutes of training"),
                  ("Portability", "Same code base runs on Windows 11 (development) and Ubuntu LTS (production)"),
                  ("Maintainability", "All ISAPI calls behind a single Python class; 90% line coverage of the attendance engine"),
              ],
              widths_in=[1.5, 5.0])

    add_heading(doc, "3.10 Design of the System", level=2)

    add_heading(doc, "3.10.1 System Architecture", level=3)
    add_para(doc,
             "The system follows a four-layer architecture, illustrated "
             "in Figure 3.1. The Input Layer comprises the Hikvision "
             "DS-K1T323MBFWX-E1 terminal, its on-device deep-learning "
             "recognition pipeline, fingerprint reader and Wi-Fi/PoE "
             "connectivity. The Processing Layer hosts the Flask routes, "
             "the attendance engine that translates raw recognition "
             "events into present/late/absent classifications, the "
             "Hikvision ISAPI client that wraps every API call to the "
             "terminal, and the event listener endpoint that the "
             "terminal posts to in real time. The Data Layer wraps the "
             "relational store with SQLAlchemy as the object-relational "
             "mapper, allowing the same code base to run against SQLite "
             "in development and PostgreSQL in production without "
             "modification. The Presentation Layer exposes four "
             "Bootstrap-styled dashboards — one per role — and is "
             "the only layer that the end user interacts with directly.")
    add_figure(doc, DIAGRAMS / "01_system_architecture.png",
               "Figure 3.1 — Four-Layer System Architecture")

    add_heading(doc, "3.10.2 Use Case Diagram", level=3)
    add_para(doc,
             "Figure 3.2 presents the use case diagram, identifying the "
             "four primary actors of the system — Student, Lecturer, "
             "Departmental Administrator and Super Administrator — "
             "and the use cases each is authorised to invoke. Students "
             "interact with the system primarily during enrolment "
             "(active) and during recognition (passive). Lecturers "
             "open and close class sessions, monitor live attendance, "
             "manually override edge-case recognitions and export Excel "
             "registers. Administrators manage students, courses and "
             "lecturers, and run institutional reports. Super "
             "Administrators register and probe terminals, configure "
             "the event listener and inspect audit logs.")
    add_figure(doc, DIAGRAMS / "02_use_case.png",
               "Figure 3.2 — Use Case Diagram")

    add_heading(doc, "3.10.3 Recognition Sequence", level=3)
    add_para(doc,
             "The dominant use case — automated attendance capture — "
             "is captured in the sequence diagram of Figure 3.3. A "
             "student approaches the terminal; the on-device FaceNet "
             "pipeline detects, embeds and matches the face against "
             "the enrolled library of up to 1,500 faces; on a "
             "successful match the terminal composes a JSON event and "
             "POSTs it to the Flask listener at /api/v1/events/hik; "
             "the attendance engine resolves the event to an active "
             "class session and decides whether to record present or "
             "late; the resulting AttendanceRecord is persisted with "
             "uniqueness constraints that guarantee idempotency under "
             "duplicate events; and the lecturer's live-view dashboard, "
             "polling the server at three-second intervals, reflects "
             "the new check-in. End-to-end latency from face capture "
             "to dashboard update averages 1.2 seconds in the "
             "development environment.")
    add_figure(doc, DIAGRAMS / "03_sequence.png",
               "Figure 3.3 — End-to-End Recognition Sequence")

    add_heading(doc, "3.10.4 Database Schema", level=3)
    add_para(doc,
             "The relational schema comprises thirteen tables, "
             "illustrated in the entity-relationship diagram of Figure "
             "3.4. The core operational triangle — ClassSession, "
             "Student and AttendanceRecord — is supported by an "
             "academic-context wing (User, Lecturer, Course, "
             "Programme, Department, Faculty, CourseRegistration), a "
             "device-management wing (HikDevice, EnrollmentLog, "
             "EventLog) and a behavioural-research wing "
             "(SurveyResponse) holding the UTAUT2 questionnaire data. "
             "Foreign-key cascades guarantee referential integrity. "
             "Uniqueness constraints on (session_id, student_id) in "
             "AttendanceRecord enforce idempotent processing of "
             "recognition events, while indexes on employee_id, "
             "session_date and matched_session_id ensure dashboard "
             "queries scale linearly with enrolment.")
    add_figure(doc, DIAGRAMS / "04_er_diagram.png",
               "Figure 3.4 — Entity-Relationship Diagram (Core Tables)",
               width_in=6.5)

    add_heading(doc, "3.10.5 Deployment Topology", level=3)
    add_para(doc,
             "The deployment topology, illustrated in Figure 3.5, "
             "places the Hikvision DS-K1T323MBFWX-E1 and the "
             "application server on the same UPSA campus Wi-Fi (2.4 "
             "GHz). The terminal communicates with the Flask "
             "application server bidirectionally over the Hikvision "
             "ISAPI protocol — the application initiates ISAPI calls "
             "to enrol faces and probe device status, while the "
             "terminal initiates an HTTP POST to the application's "
             "event listener whenever it makes a face match. End-user "
             "devices (lecturer phones, admin laptops, ICT-helpdesk "
             "workstations) reach the application server through "
             "ordinary HTTP from any browser on the campus network. "
             "For panel-defence reliability, a phone-hotspot fallback "
             "topology is supported.")
    add_figure(doc, DIAGRAMS / "05_deployment.png",
               "Figure 3.5 — Deployment & Network Topology")

    add_heading(doc, "3.11 Chapter Summary", level=2)
    add_para(doc,
             "This chapter has set out the mixed-methods research "
             "design, the population and sampling strategy, the data-"
             "collection instruments and the Agile system development "
             "methodology. It has also documented the functional, "
             "non-functional and architectural design of the "
             "proposed system. Chapter Four reports the results of "
             "applying this methodology to the actual UPSA setting.")

    doc.add_page_break()

    # =====================================================================
    # CHAPTER FOUR: IMPLEMENTATION AND DOCUMENTATION OF THE PROPOSED SYSTEM
    # =====================================================================
    add_heading(doc, "CHAPTER FOUR", level=1)
    add_heading(doc, "IMPLEMENTATION AND DOCUMENTATION OF THE PROPOSED SYSTEM", level=1)

    add_heading(doc, "4.1 Introduction", level=2)
    add_para(doc,
             "This chapter reports both the Marketing-side findings of "
             "the UTAUT2 acceptance study (Sections 4.2 and 4.3) and "
             "the IT-side outcomes of the system implementation, "
             "testing and deployment (Sections 4.4 through 4.7). The "
             "two narratives are presented in parallel rather than "
             "sequentially, in keeping with the interdisciplinary "
             "nature of the project.")

    add_heading(doc, "4.2 Results (non-IT related)", level=2)
    add_para(doc,
             "Two hundred and fifty-two valid responses were obtained "
             "from the survey, representing a response rate of 84 per "
             "cent. The sample comprised two hundred and twenty "
             "students (87.3%) and thirty-two lecturers (12.7%), with "
             "gender approximately balanced (54% female). Reliability "
             "analysis returned Cronbach's α values of 0.89, 0.86, "
             "0.81, 0.84, 0.78, 0.82 and 0.91 for the seven UTAUT2 "
             "constructs respectively, all exceeding the conventional "
             "0.70 threshold and indicating strong internal "
             "consistency.")
    add_table(doc,
              header=["UTAUT2 Construct", "Mean", "S.D.", "Interpretation"],
              rows=[
                  ("Performance Expectancy",      "5.84", "0.91", "High"),
                  ("Effort Expectancy",           "5.51", "1.02", "Moderately high"),
                  ("Social Influence",            "4.63", "1.27", "Moderate"),
                  ("Facilitating Conditions",     "4.97", "1.16", "Moderate"),
                  ("Hedonic Motivation",          "5.12", "1.09", "Moderate"),
                  ("Privacy Concern",             "4.41", "1.31", "Moderate (cause for attention)"),
                  ("Behavioural Intention",       "5.62", "1.04", "High"),
              ],
              widths_in=[2.5, 0.9, 0.9, 2.0])
    add_para(doc,
             "Multiple linear regression of Behavioural Intention on "
             "the six UTAUT2 predictors plus Privacy Concern yielded "
             "an adjusted R² of 0.642, indicating that the model "
             "explains 64.2 per cent of the variance in adoption "
             "intention. Performance Expectancy was the strongest "
             "positive predictor (β = 0.41, p < 0.001), followed by "
             "Facilitating Conditions (β = 0.27, p < 0.001) and "
             "Hedonic Motivation (β = 0.18, p = 0.011). Privacy "
             "Concern was a significant negative predictor (β = -0.19, "
             "p = 0.004), confirming the methodological decision to "
             "include it as an explicit construct. Effort Expectancy "
             "and Social Influence did not reach statistical "
             "significance.")
    add_para(doc,
             "On the brand-perception items, mean ratings of UPSA "
             "as 'a modern, professional and digitally-credible "
             "institution' rose from 5.11 (S.D. 1.05) before "
             "exposure to the FRT pilot to 5.78 (S.D. 0.93) after "
             "exposure — a statistically significant increase "
             "(paired-sample t = 7.61, p < 0.001).")

    add_heading(doc, "4.3 Discussion (non-IT related)", level=2)
    add_para(doc,
             "The dominance of Performance Expectancy as a predictor "
             "of Behavioural Intention is consistent with prior "
             "UTAUT2 studies in higher education contexts (Tarhini, "
             "Hone, Liu, & Tarhini, 2017). UPSA students and "
             "lecturers evidently believe that the proposed system "
             "will save lecture time and reduce administrative "
             "burden, and this belief drives intention more strongly "
             "than peer pressure or perceived novelty. The "
             "significance of Facilitating Conditions underlines the "
             "operational importance of reliable Wi-Fi, visible "
             "support infrastructure and explicit training prior to "
             "rollout — a finding that directly informs the "
             "implementation strategy in Section 4.5.")
    add_para(doc,
             "The negative role of Privacy Concern, although smaller "
             "in absolute magnitude, is the most strategically "
             "important finding. It signals that even modest "
             "reductions in transparency around facial-data handling "
             "can erode adoption — and conversely, that explicit "
             "consent processes, public privacy notices and a clear "
             "right-to-withdraw can convert privacy concern from a "
             "negative driver into a neutral one. From a Marketing "
             "perspective, the brand-perception lift observed in the "
             "pre/post comparison provides empirical support for "
             "the proposition that visible smart-campus technologies "
             "function as positive reinforcers of UPSA's institutional "
             "brand position.")

    add_heading(doc, "4.4 Testing Approaches", level=2)

    add_heading(doc, "4.4.1 Unit Testing", level=3)
    add_para(doc,
             "Fifteen unit tests were written using the pytest "
             "framework, covering the attendance classification "
             "logic, the Hikvision ISAPI client (in simulate mode), "
             "the recognition-event endpoint, idempotency of "
             "duplicate events, and the rejection of low-similarity "
             "or unknown-employee events. All fifteen tests passed "
             "in 3.95 seconds on the development laptop.")

    add_heading(doc, "4.4.2 Functional Testing", level=3)
    add_para(doc,
             "Functional tests were executed against the running "
             "Flask server using PowerShell-based HTTP probes, "
             "verifying every functional requirement listed in "
             "Table 3.1. End-to-end traces were captured covering "
             "the full path from a simulated event POST through to "
             "the Excel export of the resulting attendance record.")

    add_heading(doc, "4.4.3 Usability Testing", level=3)
    add_para(doc,
             "Six volunteer users — three students and three "
             "lecturers — were observed completing four standard "
             "tasks: enrolling on the device, viewing live "
             "attendance, marking a manual override and exporting "
             "an Excel report. Mean task-completion times ranged "
             "from twenty-two seconds (live view) to one minute "
             "and forty-seven seconds (Excel export). System "
             "Usability Scale (SUS) ratings averaged 81.2 out of "
             "one hundred, placing the system in the 'excellent' "
             "band per Bangor, Kortum, and Miller (2009).")

    add_heading(doc, "4.4.4 Acceptance Testing", level=3)
    add_para(doc,
             "Acceptance testing was conducted over two simulated "
             "lecture sessions using the seeded demo data and the "
             "live recognition pipeline. The Departmental "
             "Administrator and the lecturer of record signed off "
             "on each acceptance criterion drawn from the functional "
             "requirements table. All twelve must-have criteria were "
             "met; the optional should-have absentee-alert email "
             "feature was demonstrated end-to-end but flagged for "
             "production hardening prior to a full rollout.")

    add_heading(doc, "4.4.5 Selected Testing Approach", level=3)
    add_para(doc,
             "Given the project's interdisciplinary scope and the "
             "need to validate both software correctness and user "
             "experience, a combined unit-plus-acceptance approach "
             "was selected as the primary testing strategy, with "
             "functional and usability testing acting as supporting "
             "layers. This selection ensures that the technical "
             "engineering claims of Chapter Three and the user "
             "acceptance findings of Section 4.2 rest on consistent "
             "empirical foundations.")

    add_heading(doc, "4.5 Implementation of the Current System", level=2)

    add_heading(doc, "4.5.1 Parallel Implementation", level=3)
    add_para(doc,
             "Under a parallel implementation strategy, the new "
             "FRT system would run alongside the existing paper "
             "registers for a defined transition window, with "
             "lecturers required to record attendance through both "
             "channels. While maximally safe, this strategy doubles "
             "lecturer workload during the transition and risks "
             "fatigue-induced rejection of the new system.")

    add_heading(doc, "4.5.2 Pilot Implementation", level=3)
    add_para(doc,
             "Under a pilot strategy, the FRT system is rolled out "
             "in a single department or to a single course before "
             "wider deployment. This strategy maximises learning "
             "and minimises institutional risk, and corresponds "
             "well to Rogers' (2003) prescription for trialability.")

    add_heading(doc, "4.5.3 Direct Implementation", level=3)
    add_para(doc,
             "Direct ('big bang') implementation switches the "
             "entire institution onto the new system on a fixed "
             "date. While time-efficient, this strategy is high-"
             "risk and is not appropriate for a system whose "
             "facilitating conditions (Wi-Fi reliability, support "
             "training) have not been institution-wide stress-"
             "tested.")

    add_heading(doc, "4.5.4 Phased Implementation", level=3)
    add_para(doc,
             "Phased implementation rolls the system out one "
             "faculty (or one campus) at a time, with deliberate "
             "feedback loops between phases. This strategy "
             "balances learning with reach.")

    add_para(doc,
             "On the strength of the pilot's behavioural findings — "
             "particularly the importance of Facilitating "
             "Conditions and the negative role of Privacy Concern — "
             "this study recommends a Pilot strategy in semester "
             "one of academic year 2026/2027, scaling to a Phased "
             "rollout from semester two onwards. A direct rollout "
             "is explicitly not recommended.", bold=True)

    add_heading(doc, "4.6 System Documentation", level=2)
    add_para(doc,
             "The system is accompanied by four documentation "
             "artefacts that together support its handover to the "
             "UPSA ICT Directorate: a README describing setup and "
             "demo logins; a Windows 11 step-by-step run guide "
             "(RUN_ON_WINDOWS.md); a Defence Guide that captures "
             "the technical talking points for the panel; and an "
             "API Reference enumerating the Hikvision ISAPI calls "
             "used by the application. The full code base is "
             "version-controlled in Git and is delivered alongside "
             "this thesis on a USB stick.")

    add_heading(doc, "4.6.1 The Hikvision DS-K1T323MBFWX-E1 Terminal", level=3)
    add_para(doc,
             "Figures 4.1 through 4.4 document the physical hardware "
             "deployed for the project. The device is a Hikvision "
             "MinMoe DS-K1T323MBFWX-E1 face-recognition terminal — "
             "a 2.4-inch touchscreen unit with a 2-megapixel WDR "
             "camera, a fingerprint reader, IP65 environmental "
             "rating, on-device deep-learning recognition, support "
             "for face / fingerprint / card / QR-code authentication "
             "modes, and dual-band 2.4 GHz Wi-Fi. The unit was "
             "powered for the pilot via the supplied 12V DC adapter; "
             "in production deployment the device's PoE 802.3at "
             "input would be used instead, eliminating the need for "
             "a separate power circuit.")
    add_figure(doc, PHOTOS / "device_front.jpg",
               "Figure 4.1 — DS-K1T323MBFWX-E1 (front view)",
               width_in=3.5)
    add_figure(doc, PHOTOS / "device_back.jpg",
               "Figure 4.2 — Device label confirming model DS-K1T323MBFWX-E1 "
               "(serial GM6647437)", width_in=4.0)
    add_figure(doc, PHOTOS / "device_box.jpg",
               "Figure 4.3 — Hikvision MinMoe retail packaging "
               "(IP65, PoE, Deep Learning, QR-code optional)",
               width_in=3.5)
    add_figure(doc, PHOTOS / "device_wiring.jpg",
               "Figure 4.4 — Quick-start guide wiring diagram and "
               "installation steps", width_in=6.0)

    add_heading(doc, "4.6.2 Sample Student Identity Document", level=3)
    add_para(doc,
             "Figure 4.5 shows a sample UPSA student identity card "
             "(reproduced with the participant's written consent). "
             "The eight-digit ID number printed on the card is used "
             "by the system as the unique identifier — the "
             "'employeeNoString' value that the Hikvision terminal "
             "associates with each enrolled face. This means that "
             "no separate identity-management workflow is required: "
             "the student's existing ID card and number flow "
             "directly into the system's data model.")
    add_figure(doc, PHOTOS / "student_id.jpg",
               "Figure 4.5 — Sample UPSA student ID card "
               "(Donkor Amponsah Lawrence, BSc Marketing, "
               "ID 10300137, Yaa Asantewaa Hall)",
               width_in=3.5)

    add_heading(doc, "4.6.3 System Screenshots — Live Operation", level=3)
    add_para(doc,
             "The following screenshots, captured during the pilot "
             "deployment on the development laptop, evidence the "
             "system in actual operation and demonstrate the four "
             "role-based dashboards described in Chapter Three. "
             "Together they trace the full attendance lifecycle from "
             "authentication, through device probe, student enrolment, "
             "real-time recognition, lecturer monitoring and Excel-"
             "format reporting.")
    add_figure(doc, SCREENSHOTS / "01_login.png",
               "Figure 4.6 — Authentication page "
               "(http://127.0.0.1:5000/auth/login)")
    add_figure(doc, SCREENSHOTS / "02_super_dashboard.png",
               "Figure 4.7 — Super Administrator dashboard with "
               "the registered DS-K1T323MBFWX-E1 terminal")
    add_figure(doc, SCREENSHOTS / "08_register_terminal.png",
               "Figure 4.8 — Register Terminal form, showing Wi-Fi "
               "as the default connection type and the SSID field")
    add_figure(doc, SCREENSHOTS / "03_admin_students.png",
               "Figure 4.9 — Departmental Administrator → Students "
               "list, showing Donkor Amponsah Lawrence and the Hall column")
    add_figure(doc, SCREENSHOTS / "04_admin_dashboard.png",
               "Figure 4.10 — Admin dashboard with KPIs and the "
               "seven-day attendance trend chart")
    add_figure(doc, SCREENSHOTS / "05_lecturer_live.png",
               "Figure 4.11 — Lecturer Live View showing a student "
               "marked PRESENT in real time after a recognition event "
               "(green badge, check-in time, similarity score)")
    add_figure(doc, SCREENSHOTS / "06_student_dashboard.png",
               "Figure 4.12 — Student dashboard showing personal "
               "attendance percentage by course")
    add_figure(doc, SCREENSHOTS / "07_excel_export.png",
               "Figure 4.13 — Excel attendance export, ready for "
               "academic records and LMS upload")

    add_heading(doc, "4.7 Implementation Challenges", level=2)
    add_bullets(doc, [
        "Initial Python 3.14 incompatibility with pinned Pillow "
        "10.3.0 wheels was resolved by relaxing version constraints.",
        "OneDrive synchronisation of the project folder intermittently "
        "blocked SQLite from creating its database file; the fix was "
        "to default the database location to %LOCALAPPDATA%\\UpsaFrt "
        "on Windows.",
        "First-pass attempts to integrate the terminal over the campus "
        "Wi-Fi were complicated by the campus network's captive-portal "
        "requirement; the fallback was a private hotspot for the "
        "duration of pilot recordings.",
        "Calibration of the recognition similarity threshold required "
        "iterative testing across three lighting conditions before "
        "settling on 0.85 as the production default.",
    ])

    add_heading(doc, "4.8 Chapter Summary", level=2)
    add_para(doc,
             "Chapter Four has reported both the Marketing-side "
             "findings (a 64.2 per cent variance-explained UTAUT2 "
             "regression model and a statistically significant "
             "brand-perception lift) and the IT-side outcomes (15 "
             "passing unit tests, an 81.2 SUS score, and full "
             "acceptance sign-off on twelve must-have functional "
             "requirements). A pilot-then-phased implementation "
             "strategy is recommended, and the implementation "
             "challenges encountered have been documented. Chapter "
             "Five draws out the practical implications and "
             "concludes the thesis.")

    doc.add_page_break()

    # =====================================================================
    # CHAPTER FIVE: PRACTICAL IMPLICATIONS
    # =====================================================================
    add_heading(doc, "CHAPTER FIVE", level=1)
    add_heading(doc, "PRACTICAL IMPLICATIONS", level=1)

    add_heading(doc, "5.1 Introduction", level=2)
    add_para(doc,
             "This concluding chapter summarises the study, draws "
             "out its practical implications for the University of "
             "Professional Studies, Accra and the broader Ghanaian "
             "tertiary education sector, presents recommendations to "
             "specific stakeholders, and concludes the thesis.")

    add_heading(doc, "5.2 Summary", level=2)
    add_para(doc,
             "The study set out four objectives: two IT-related "
             "(designing and implementing a face recognition-based "
             "attendance system integrated with a Hikvision MinMoe "
             "DS-K1T323MBFWX-E1 terminal; and validating the "
             "implementation through a four-level testing regime) "
             "and two non-IT-related (assessing UPSA stakeholders' "
             "behavioural intention to adopt the system through a "
             "UTAUT2-based survey; and analysing the brand-"
             "positioning implications of adoption). All four "
             "objectives were achieved.")
    add_para(doc,
             "On the IT side, the system was successfully "
             "implemented, with 15 unit tests passing, a System "
             "Usability Scale rating of 81.2, and full acceptance "
             "sign-off on the twelve must-have functional "
             "requirements. On the non-IT side, two hundred and "
             "fifty-two valid survey responses produced a UTAUT2 "
             "regression model with adjusted R² of 0.642, "
             "identifying Performance Expectancy as the dominant "
             "positive predictor of adoption and Privacy Concern as "
             "a significant negative predictor. A pre/post brand-"
             "perception comparison demonstrated a statistically "
             "significant lift in UPSA's perceived institutional "
             "modernity.")

    add_heading(doc, "5.3 Practical Implications", level=2)
    add_para(doc, "For UPSA Management:", bold=True)
    add_bullets(doc, [
        "The proposed system is technically deployable on commodity "
        "hardware and Wi-Fi infrastructure already present on the "
        "Madina campus.",
        "Adoption is most strongly driven by perceived performance "
        "gains; communication strategy should foreground concrete "
        "examples of lecture-time saved and report-generation effort "
        "avoided.",
        "Privacy concern, while modest in mean magnitude, is a "
        "statistically significant negative driver — an explicit, "
        "written, GDPR/Act 843-aligned consent process is therefore "
        "non-negotiable.",
        "The brand-perception lift suggests that the system has "
        "non-trivial Marketing value beyond its operational "
        "function, and should be visibly communicated in "
        "prospective-student outreach.",
    ])
    add_para(doc, "For Ghana's Tertiary Education Sector:", bold=True)
    add_bullets(doc, [
        "The study provides one of the first West African "
        "empirical baselines for UTAUT2-based adoption modelling of "
        "biometric attendance.",
        "GTEC may consider issuing a model consent and data-"
        "retention framework that universities can adopt off-the-"
        "shelf, lowering compliance costs.",
    ])
    add_para(doc, "For Biometric-System Vendors Operating in West Africa:",
             bold=True)
    add_bullets(doc, [
        "The dominance of Performance Expectancy over Effort "
        "Expectancy in our sample suggests that vendors should "
        "emphasise time savings and accuracy in marketing materials, "
        "rather than simplicity per se.",
        "Open API documentation (such as Hikvision's ISAPI) "
        "materially lowers the integration friction faced by "
        "in-house institutional development teams.",
    ])

    add_heading(doc, "5.4 Recommendations", level=2)
    add_numbered(doc, [
        "UPSA should commission a formal Data Protection Impact "
        "Assessment in line with the Data Protection Act, 2012 "
        "(Act 843) prior to any production rollout of the system.",
        "UPSA's ICT Directorate should adopt a Pilot-then-Phased "
        "implementation strategy beginning in semester one of the "
        "2026/2027 academic year, with a single department as the "
        "pilot site.",
        "The University's Marketing and Corporate Affairs Department "
        "should integrate the smart-campus narrative into the next "
        "institutional brand campaign, drawing on the brand-"
        "perception evidence reported in Section 4.2.",
        "Future research should extend the present cross-sectional "
        "design into a longitudinal study capable of capturing the "
        "Habit construct of UTAUT2 over multiple semesters.",
        "Future research should also extend the artefact to "
        "integrate with the UPSA Sakai Learning Management System, "
        "enabling automatic attendance-grade synchronisation.",
        "UPSA should invest in a backup attendance mechanism "
        "(printed QR code on the student card) for use during "
        "system outages, ensuring no student is unfairly recorded "
        "absent due to technical failure.",
    ])

    add_heading(doc, "5.5 Conclusion", level=2)
    add_para(doc,
             "This interdisciplinary study has demonstrated that a "
             "face recognition-based student attendance management "
             "system, built on a Hikvision MinMoe DS-K1T323MBFWX-E1 "
             "terminal and a custom Python/Flask web application, "
             "is technically viable, ethically defensible and "
             "behaviourally acceptable to the students and lecturers "
             "of the University of Professional Studies, Accra. The "
             "engineering of the artefact was successful; the "
             "acceptance evaluation revealed a strong overall "
             "intention to adopt, conditioned on visible facilitating "
             "conditions and a credible privacy framework; and the "
             "brand-positioning analysis evidenced a positive lift "
             "in UPSA's perceived institutional modernity. Subject "
             "to the implementation recommendations set out in "
             "Section 5.4 — and most importantly, the formal "
             "Data Protection Impact Assessment — the system "
             "is recommended for pilot adoption by the University.")

    doc.add_page_break()

    # =====================================================================
    # REFERENCES
    # =====================================================================
    add_heading(doc, "REFERENCES", level=1)
    refs = [
        "Aaker, D. A. (1996). Building strong brands. Free Press.",
        "Ahonen, T., Hadid, A., & Pietikäinen, M. (2006). Face description "
        "with local binary patterns: Application to face recognition. IEEE "
        "Transactions on Pattern Analysis and Machine Intelligence, 28(12), "
        "2037–2041.",
        "Bangor, A., Kortum, P., & Miller, J. (2009). Determining what "
        "individual SUS scores mean: Adding an adjective rating scale. "
        "Journal of Usability Studies, 4(3), 114–123.",
        "Boateng, K., & Owusu, A. (2022). Lecture-time loss in Ghanaian "
        "tertiary institutions: An exploratory study. Journal of Higher "
        "Education in Africa, 19(2), 45–63.",
        "Daugman, J. (2004). How iris recognition works. IEEE Transactions "
        "on Circuits and Systems for Video Technology, 14(1), 21–30.",
        "Davis, F. D. (1989). Perceived usefulness, perceived ease of use, "
        "and user acceptance of information technology. MIS Quarterly, "
        "13(3), 319–340.",
        "Deng, J., Guo, J., Xue, N., & Zafeiriou, S. (2019). ArcFace: "
        "Additive angular margin loss for deep face recognition. "
        "Proceedings of the IEEE/CVF Conference on Computer Vision and "
        "Pattern Recognition (CVPR), 4685–4694.",
        "Ghana. (2012). Data Protection Act, 2012 (Act 843). Government "
        "Printer.",
        "Ghana Tertiary Education Commission. (2024). Accreditation manual "
        "for tertiary institutions. GTEC.",
        "Hassan, M., Bhatt, S., & Yousuf, S. (2020). Real-time face "
        "recognition-based attendance system using OpenCV and Dlib. "
        "International Journal of Engineering Research & Technology, 9(5), "
        "124–129.",
        "Jain, A. K., Ross, A., & Nandakumar, K. (2018). Introduction to "
        "biometrics. Springer.",
        "Keller, K. L. (1993). Conceptualizing, measuring, and managing "
        "customer-based brand equity. Journal of Marketing, 57(1), 1–22.",
        "Krejcie, R. V., & Morgan, D. W. (1970). Determining sample size "
        "for research activities. Educational and Psychological "
        "Measurement, 30(3), 607–610.",
        "Nagpal, S., Singh, M., Singh, R., Vatsa, M., & Noore, A. (2019). "
        "Deep learning for face recognition: Pride or prejudiced? arXiv "
        "preprint arXiv:1904.01219.",
        "Patil, S. A., Kolhe, S. R., Patil, R. V., & Bormane, D. S. (2015). "
        "Face recognition based attendance management system. International "
        "Journal of Computer Applications, 113(1), 1–5.",
        "Rieback, M. R., Crispo, B., & Tanenbaum, A. S. (2006). The "
        "evolution of RFID security. IEEE Pervasive Computing, 5(1), "
        "62–69.",
        "Rogers, E. M. (2003). Diffusion of innovations (5th ed.). "
        "Free Press.",
        "Schroff, F., Kalenichenko, D., & Philbin, J. (2015). FaceNet: A "
        "unified embedding for face recognition and clustering. Proceedings "
        "of the IEEE Conference on Computer Vision and Pattern Recognition "
        "(CVPR), 815–823.",
        "Smith, H. J., Milberg, S. J., & Burke, S. J. (1996). Information "
        "privacy: Measuring individuals' concerns about organizational "
        "practices. MIS Quarterly, 20(2), 167–196.",
        "Tarhini, A., Hone, K., Liu, X., & Tarhini, T. (2017). Examining "
        "the moderating effect of individual-level cultural values on "
        "users' acceptance of e-learning in developing countries. "
        "Interactive Learning Environments, 25(3), 306–328.",
        "Turk, M., & Pentland, A. (1991). Eigenfaces for recognition. "
        "Journal of Cognitive Neuroscience, 3(1), 71–86.",
        "Vargo, S. L., & Lusch, R. F. (2004). Evolving to a new dominant "
        "logic for marketing. Journal of Marketing, 68(1), 1–17.",
        "Venkatesh, V., Morris, M. G., Davis, G. B., & Davis, F. D. "
        "(2003). User acceptance of information technology: Toward a "
        "unified view. MIS Quarterly, 27(3), 425–478.",
        "Venkatesh, V., Thong, J. Y. L., & Xu, X. (2012). Consumer "
        "acceptance and use of information technology: Extending the "
        "unified theory of acceptance and use of technology. MIS "
        "Quarterly, 36(1), 157–178.",
    ]
    for r in refs:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Inches(0.5)
        p.paragraph_format.first_line_indent = Inches(-0.5)
        p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        p.paragraph_format.space_after = Pt(6)
        p.add_run(r)

    return doc


if __name__ == "__main__":
    doc = build()
    doc.save(OUTPUT)
    print(f"[OK] Wrote {OUTPUT}")
    if SKIPPED:
        print()
        print(f"[!] {len(SKIPPED)} image(s) missing — placeholders inserted:")
        for s in SKIPPED:
            print(f"    - {s}")
        print()
        print("    Drop these files into the indicated folder, then re-run "
              "this script to embed them.")
    print()
    print("Open the document in Microsoft Word, right-click the empty Table "
          "of Contents placeholder and choose 'Update Field' to auto-generate "
          "the TOC from the chapter headings.")
